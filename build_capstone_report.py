from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


OUT = Path("outputs")
OUT.mkdir(exist_ok=True)
DOCX_PATH = OUT / "ResolveNow_AI_Capstone_Report_Priyadharshini.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
GRAY = RGBColor(89, 89, 89)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
NO_BREAK_AFTER = {
    "1.1 Introduction",
    "1.2 Background and Motivation",
    "1.3 Problem Statement",
    "2.2 Scope of the Project",
    "2.3 Literature Review - Incident Classification",
    "2.4 Literature Review - NLP and Knowledge Extraction",
    "2.5 Literature Review - GNN and RAG",
    "3.1 Data Preprocessing",
    "3.2 Exploratory Data Analysis",
    "4.2 FastAPI Workflow",
    "4.3 Model Pipeline",
    "4.4 BERT and Semantic Understanding",
    "4.5 RAG and FAISS Retrieval",
    "5.2 Dashboard Implementation",
    "6.2 Model Evaluation Outputs",
    "7.1 Data and Modeling Challenges",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(9.5)
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2EC")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_bottom_border(paragraph, color="2E74B5", size="8"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.10

    if "Report Caption" not in styles:
        cap = styles.add_style("Report Caption", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = "Calibri"
        cap.font.size = Pt(9)
        cap.font.italic = True
        cap.font.color.rgb = GRAY
        cap.paragraph_format.space_after = Pt(6)

    if "Callout" not in styles:
        call = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
        call.font.name = "Calibri"
        call.font.size = Pt(10.5)
        call.paragraph_format.left_indent = Inches(0.15)
        call.paragraph_format.right_indent = Inches(0.15)
        call.paragraph_format.space_before = Pt(4)
        call.paragraph_format.space_after = Pt(8)
        call.paragraph_format.line_spacing = 1.12

    header = section.header.paragraphs[0]
    header.text = "ResolveNow AI Capstone Project Report"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = GRAY
    add_bottom_border(header, "D9E2EC", "4")

    footer = section.footer.paragraphs[0]
    footer.add_run("Priyadharshini J - 2514018").font.size = Pt(9)
    add_page_number(footer)


def p(doc, text="", style=None, align=None, bold=False, italic=False, color=None, size=None):
    para = doc.add_paragraph(style=style)
    if align:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.first_line_indent = Inches(-0.25)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.167
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.first_line_indent = Inches(-0.25)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.167
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    r = para.add_run(title + ": ")
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r2 = para.add_run(text)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10.5)
    set_table_borders(table)
    p(doc, "")


def add_table(doc, headers, rows, widths=None, caption=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    if widths:
        for i, w in enumerate(widths):
            table.columns[i].width = Inches(w)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_FILL)
        set_cell_text(hdr[i], h, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    if caption:
        p(doc, caption, style="Report Caption", align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        p(doc, "")
    return table


def add_image(doc, path, caption, width=6.1):
    img = Path(path)
    if img.exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(img), width=Inches(width))
        p(doc, caption, style="Report Caption", align=WD_ALIGN_PARAGRAPH.CENTER)


def section_page(doc, heading, paragraphs=None, bullets=None, table=None, image=None, callout=None):
    doc.add_heading(heading, level=1 if heading.startswith("CHAPTER") else 2)
    if callout:
        add_callout(doc, callout[0], callout[1])
    for text in paragraphs or []:
        p(doc, text)
    for text in bullets or []:
        bullet(doc, text)
    if table:
        add_table(doc, table["headers"], table["rows"], table.get("widths"), table.get("caption"))
    if image:
        add_image(doc, image["path"], image["caption"], image.get("width", 6.1))
    if heading not in NO_BREAK_AFTER:
        doc.add_page_break()


def cover(doc):
    for _ in range(3):
        p(doc, "")
    p(doc, "ResolveNow AI", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=BLUE, size=28)
    p(doc, "Autonomous ITSM Incident Intelligence System", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=DARK_BLUE, size=16)
    p(doc, "Capstone Project Report", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    for _ in range(2):
        p(doc, "")
    p(doc, "Submitted by", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    p(doc, "Priyadharshini J", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    p(doc, "Register Number: 2514018", align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "PGDDSBA", align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(2):
        p(doc, "")
    p(doc, "Under the guidance of", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    p(doc, "Dr. Nataraj", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    p(doc, "Thiagarajar School of Management", align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "17 June 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def front_matter(doc):
    section_page(
        doc,
        "CERTIFICATE",
        [
            "This is to certify that the capstone project report titled 'ResolveNow AI - Autonomous ITSM Incident Intelligence System' is a record of work carried out by Priyadharshini J, Register Number 2514018, as part of the PGDDSBA capstone project requirement.",
            "The project demonstrates an applied analytics and AI solution for IT Service Management by combining ticket classification, known error retrieval, priority-based routing, email automation, SLA monitoring, and dashboard-based ticket lifecycle management.",
            "The work presented in this report has been prepared using the final project presentation, source code implementation, generated model artifacts, dataset analysis, and functional dashboard/backend components available in the ResolveNow AI project workspace.",
        ],
        callout=("Certification note", "The report is intended for academic evaluation and describes the working portions of the ResolveNow AI prototype."),
    )
    section_page(
        doc,
        "DECLARATION",
        [
            "I, Priyadharshini J, hereby declare that this capstone project report titled 'ResolveNow AI - Autonomous ITSM Incident Intelligence System' is based on my own project implementation, analysis, and presentation work.",
            "The project applies artificial intelligence and automation concepts to an ITSM incident management use case. The report combines the final presentation content with the working codebase, dataset preparation logic, backend API implementation, dashboard behavior, email workflow, KEDB construction, RAG retrieval, GNN prototype, and voice chatbot feature.",
            "The contents of this report are prepared for academic submission and are not copied from any previous report. External ideas discussed in the literature review are acknowledged as conceptual references.",
        ],
    )
    section_page(
        doc,
        "ACKNOWLEDGEMENT",
        [
            "I express my sincere gratitude to Dr. Nataraj for the guidance, academic support, and feedback provided throughout the capstone project. The direction helped shape the project from an initial ITSM automation idea into a complete applied AI prototype.",
            "I thank Thiagarajar School of Management for providing the learning environment, project structure, and opportunity to apply data science and business analytics concepts to a practical enterprise problem.",
            "I also acknowledge the value of the ITSM incident dataset used for experimentation and the open-source tools and libraries that supported the project, including FastAPI, SQLite, scikit-learn, FAISS, SentenceTransformer, PyTorch Geometric, Neo4j, and Python visualization libraries.",
        ],
    )
    section_page(
        doc,
        "ABSTRACT",
        [
            "ResolveNow AI is an autonomous IT Service Management incident intelligence platform designed to reduce manual triage effort, accelerate resolution, and improve visibility across the ticket lifecycle. The system addresses the common operational challenge where support teams spend a large proportion of incident response time searching knowledge bases, routing tickets manually, and repeating known resolutions.",
            "The project combines a live dashboard, FastAPI backend, machine learning classifiers, Known Error Database construction, RAG-based solution retrieval, GraphSAGE knowledge graph reasoning prototype, SLA monitoring, automated email workflows, and a Sarvam API based voice chatbot for ticket status tracking.",
            "The working implementation allows users to create tickets, apply AI analysis, route by priority, retrieve known solutions, send solution or escalation emails, store tickets in SQLite, monitor dashboards and analytics, and query ticket status through voice. Low and medium priority known issues are handled through solution-email and customer confirmation workflows, while high and critical tickets are escalated to agents with solution suggestions.",
        ],
        callout=("Abstract summary", "ResolveNow AI demonstrates an end-to-end AI-enabled ITSM workflow from ticket creation to routing, communication, tracking, and closure."),
    )


def contents(doc):
    doc.add_heading("TABLE OF CONTENTS", level=1)
    items = [
        "Certificate",
        "Declaration",
        "Acknowledgement",
        "Abstract",
        "Chapter 1 - Introduction",
        "1.1 Introduction",
        "1.2 Background and Motivation",
        "1.3 Problem Statement",
        "1.4 Need for an AI-Enabled ITSM Framework",
        "Chapter 2 - Objectives and Literature Review",
        "Chapter 3 - Data Source and Representation",
        "Chapter 4 - Methodology",
        "Chapter 5 - Core Strategy and Implementation",
        "Chapter 6 - Actionable Outcomes",
        "Chapter 7 - Challenges",
        "Chapter 8 - Conclusions and Recommendations",
        "Appendix A - Sample API and Data Structures",
        "Appendix B - Presentation and Demo Notes",
    ]
    for item in items:
        bullet(doc, item)
    doc.add_page_break()


def build_report():
    doc = Document()
    configure_document(doc)
    cover(doc)
    front_matter(doc)
    contents(doc)

    section_page(
        doc,
        "CHAPTER 1 - INTRODUCTION",
        [
            "ResolveNow AI is an autonomous IT Service Management platform developed to demonstrate how artificial intelligence can reduce manual effort in incident management. In a conventional ITSM process, support teams receive tickets from email, phone, chat, or portal channels, classify them manually, search past solutions, assign them to the right group, and update the requester after resolution. This workflow is repetitive, time consuming, and vulnerable to delay.",
            "The project focuses on converting this reactive process into an intelligent workflow. A user can create a ticket through the dashboard, the backend analyzes the issue, the system predicts priority and topic, checks whether the issue is already known in the KEDB, routes it based on priority, sends email communication, and stores the complete ticket record in the database.",
            "The core idea is not to replace human support teams completely. Instead, ResolveNow AI automates repetitive known-issue handling and improves the quality of escalation for high-risk incidents. Human agents receive better context and suggested solutions, while users receive faster responses for low and medium priority known issues.",
        ],
        callout=("Project position", "ResolveNow AI acts as an applied AI layer on top of ITSM operations, combining prediction, retrieval, routing, communication, and monitoring."),
    )
    section_page(
        doc,
        "1.1 Introduction",
        [
            "IT incident management is a business-critical function because downtime affects productivity, customer experience, SLA compliance, and operational cost. The ResolveNow AI project was designed around this practical business requirement. The system was built as a working prototype with a FastAPI backend, an HTML/CSS/JavaScript dashboard, SQLite persistence, trained ML artifacts, KEDB data, email services, SLA logic, and a voice chatbot.",
            "The dashboard provides the user-facing control center. It contains pages for overview metrics, all tickets, new ticket creation, KEDB viewing, and analytics. The backend contains the intelligence and workflow logic. It exposes REST APIs for ticket creation, prediction, ticket retrieval, KEDB lookup, analytics, SLA breach monitoring, ticket resolution, and voice status query.",
            "The final project output is a local web application available through the FastAPI server and dashboard route. It demonstrates the complete ticket lifecycle from creation to AI analysis, assignment, communication, and status tracking.",
        ],
    )
    section_page(
        doc,
        "1.2 Background and Motivation",
        [
            "The motivation for ResolveNow AI comes from the delay caused by manual knowledge-base searching and ticket routing. In many IT operations teams, a large share of response and restoration time is spent identifying the issue type, checking whether a similar issue has occurred previously, locating a reliable fix, and deciding which team should own the ticket.",
            "The final presentation identifies that operations teams may lose 60 to 70 percent of incident response or restoration time to manual searching, misrouting, and reactive workflows. This motivates an AI-assisted system that can interpret ticket context, find previous resolutions, and trigger the next action automatically.",
            "The project also reflects the business need to reduce Mean Time to Resolve, reduce avoidable resource effort, reduce SLA penalties, and improve user communication. These business outcomes are translated into technical capabilities: automated classification, vector search, decision logic, email automation, database audit trail, dashboard visibility, and chatbot assistance.",
        ],
    )
    section_page(
        doc,
        "1.3 Problem Statement",
        [
            "The problem addressed by the project is that traditional ITSM KEDB tools are usually reactive text search systems. They help only after a person manually searches the knowledge base. They do not automatically classify the ticket, decide the priority path, route the case, send solution steps, or close the ticket based on user confirmation.",
            "High priority and low priority tickets are often handled through the same manual workflow, even though they require different treatment. High and critical tickets need rapid human attention, while low and medium known issues can often be resolved by providing the correct solution steps to the user. Without automation, both categories consume support capacity.",
            "ResolveNow AI solves this problem by combining model-driven prediction, retrieval from historical solutions, rule-based routing, and automated communication. The system creates a structured decision layer that separates human-required incidents from automatable known issues.",
        ],
        bullets=[
            "Manual triage increases response delay and creates inconsistent routing.",
            "Reactive KEDB search does not automatically surface the right solution.",
            "High priority tickets require escalation, not the same workflow as low priority tickets.",
            "Users need faster status visibility and easier ticket tracking.",
        ],
    )
    section_page(
        doc,
        "1.4 Need for an AI-Enabled ITSM Framework",
        [
            "An AI-enabled ITSM framework is required because incident data contains repeated patterns. Topics such as network issues, access requests, hardware failures, software bugs, and general inquiries appear frequently. When these patterns are captured in a KEDB and represented through embeddings or graph relationships, the system can reuse past knowledge.",
            "ResolveNow AI introduces intelligence at multiple levels. The ML classifier predicts topic and priority confidence. RAG retrieves the most relevant KEDB solution. The GNN prototype demonstrates how graph relationships between tickets and resolutions can support SLA risk or resolution-path reasoning. SHAP provides explainability for model behavior and feature importance.",
            "The framework is also action-oriented. It does not stop at prediction. It sends emails, changes ticket status, stores records, updates dashboard metrics, and enables the user to confirm resolution. This makes the project a complete workflow automation prototype rather than only a model experiment.",
        ],
    )

    section_page(
        doc,
        "CHAPTER 2 - OBJECTIVES AND LITERATURE REVIEW",
        [
            "The project objectives were derived from the need to reduce manual incident management effort and create a more intelligent ticket lifecycle. The implementation combines several AI and software engineering components so that the final output works as an integrated application.",
            "The literature review focuses on three technical directions that influenced the project: machine learning for incident triage, NLP for extracting meaning from incident text, and retrieval-plus-graph reasoning for using historical knowledge more effectively.",
        ],
    )
    section_page(
        doc,
        "2.1 Project Objectives",
        [
            "The main objective of ResolveNow AI is to automate the repetitive parts of IT incident management while preserving human control for high-risk tickets. The solution was designed to classify incoming tickets, retrieve known fixes, route based on priority, communicate with users and agents, and provide lifecycle visibility.",
            "The system objectives are directly reflected in the working modules. The dashboard creates tickets and shows status. FastAPI accepts requests and runs the decision engine. The trained models predict priority and topic. The KEDB stores known solutions. The email service sends solution and escalation messages. The SQLite database keeps the audit trail.",
        ],
        bullets=[
            "Classify ticket topic and priority using trained ML artifacts and text features.",
            "Match tickets to known errors using KEDB lookup and RAG/FAISS retrieval concepts.",
            "Apply graph reasoning through a GraphSAGE prototype for SLA or resolution pattern learning.",
            "Route tickets according to priority and known/unknown issue status.",
            "Send intelligent emails to users and agents and support customer confirmation.",
            "Provide dashboard analytics, SLA monitoring, and voice-based ticket status tracking.",
        ],
    )
    section_page(
        doc,
        "2.2 Scope of the Project",
        [
            "The scope of the project covers a complete prototype, not a production ITSM product. It demonstrates the core logic required for autonomous incident management using local files, trained models, a SQLite database, and a local FastAPI service.",
            "The working scope includes dashboard pages, ticket creation, prediction, KEDB matching, email templates, SLA deadline calculation, breached/warning SLA APIs, ticket detail display, analytics summary, manual resolution/closure actions, Sarvam speech-to-text integration, and CSV report download.",
            "The scope does not include full enterprise deployment, user authentication, enterprise SSO, integration with ServiceNow/Jira in production, or continuous retraining automation. These are identified as future work.",
        ],
        table={
            "headers": ["In Scope", "Out of Scope / Future Work"],
            "rows": [
                ["Dashboard ticket lifecycle", "Enterprise SSO and role-based access"],
                ["FastAPI backend and SQLite persistence", "Production-scale database clustering"],
                ["ML prediction and KEDB lookup", "Automated continuous retraining pipeline"],
                ["Email workflow and confirmation links", "Native ITSM tool integration through MCP"],
                ["Sarvam STT chatbot for status query", "Full multilingual conversational support"],
            ],
            "widths": [3.1, 3.1],
            "caption": "Table 2.1 Scope boundary of the ResolveNow AI prototype",
        },
    )
    section_page(
        doc,
        "2.3 Literature Review - Incident Classification",
        [
            "DeepTriage by Microsoft Azure is referenced in the final presentation as an incident categorization approach that used ensemble models and deep networks for cloud incidents. The inference for this project is that machine learning based classification improves triage consistency when compared with manual routing.",
            "ResolveNow AI adopts this inference by training a priority classifier and topic classifier using historical ticket features. The project uses TF-IDF text representation combined with encoded categorical fields for the working classifier. The final presentation frames BERT-style semantic understanding as the conceptual direction for topic classification.",
            "The practical lesson from incident classification literature is that triage accuracy depends on combining text content with operational context. ResolveNow AI therefore includes topic, source, product group, support level, and agent group information rather than relying only on raw descriptions.",
        ],
    )
    section_page(
        doc,
        "2.4 Literature Review - NLP and Knowledge Extraction",
        [
            "SoftNER by Microsoft is referenced as an NLP approach for extracting structured knowledge such as system components and error codes from incident reports. The inference is that incident text contains useful signals that can improve routing and resolution when converted into structured features.",
            "ResolveNow AI applies this principle by creating a combined text feature from topic, source, product group, and support level. The RAG component also converts KEDB entries into embeddings using SentenceTransformer so that semantic similarity can be used for retrieval.",
            "In the report, BERT is explained as the semantic representation layer. In the working code, this is represented through SentenceTransformer embeddings for RAG and TF-IDF plus encoded features for the deployed classification artifacts.",
        ],
    )
    section_page(
        doc,
        "2.5 Literature Review - GNN and RAG",
        [
            "The GNN-RAG reference in the presentation motivates the combination of graph reasoning and retrieval. Retrieval is good at finding similar past knowledge, while graph reasoning can capture relationships among tickets, resolutions, priorities, and SLA outcomes.",
            "ResolveNow AI implements this pattern conceptually through two modules. The RAG module builds and queries a FAISS index over KEDB entries. The GNN prototype exports ticket-resolution relationships from Neo4j into a PyTorch Geometric structure and trains a GraphSAGE model for SLA prediction.",
            "This combination is important because ITSM problems are relational. A ticket is linked to a topic, priority, support group, resolution, SLA outcome, and requester. Representing these relationships as a graph provides a path for future advanced reasoning.",
        ],
    )

    section_page(
        doc,
        "CHAPTER 3 - DATA SOURCE AND REPRESENTATION",
        [
            "The project uses a Kaggle ITSM ticket dataset described in the final presentation as a real-world Gulf region enterprise operation dataset. The main file used for training and analysis is data/tickets/your_kaggle_file.csv.",
            "The dataset contains 100,000 incident records. Important columns include Status, Ticket ID, Priority, Source, Topic, Agent Group, Agent Name, Created time, SLA fields, Resolution time, Close time, Agent interactions, Survey results, Product group, Support Level, and Country.",
        ],
        table={
            "headers": ["Dataset Attribute", "Value Used in Project"],
            "rows": [
                ["Source", "Kaggle ITSM ticket dataset"],
                ["Total records", "100,000 incident tickets"],
                ["Main target fields", "Priority, Topic, Status, SLA fields"],
                ["Operational fields", "Source, Agent Group, Product group, Support Level, Country"],
                ["Time fields", "Created time, First response time, Resolution time, Close time"],
            ],
            "widths": [2.2, 4.0],
            "caption": "Table 3.1 Dataset overview",
        },
    )
    section_page(
        doc,
        "3.1 Data Preprocessing",
        [
            "The preprocessing stage cleans the dataset and prepares model-ready features. Rows with missing Priority, Topic, or Status values are dropped because these fields are central to classification and workflow decisions. Other missing categorical fields are filled with defaults such as Unknown or Neutral.",
            "A combined text feature is created by concatenating Topic, Source, Product group, and Support Level. This allows the model to learn from the issue category as well as operational context. Label encoders are then created for priority, topic, agent group, status, source, and product group.",
            "The preprocessing also supports KEDB generation by filtering resolved and closed tickets. These tickets are treated as historical cases from which known issue categories and suggested solutions can be derived.",
        ],
        bullets=[
            "Dropped records missing Priority, Topic, or Status.",
            "Filled optional categorical fields with Unknown or Neutral.",
            "Created a combined text feature for ML input.",
            "Encoded categorical fields using LabelEncoder.",
            "Used resolved and closed tickets to build KEDB categories.",
        ],
    )
    section_page(
        doc,
        "3.2 Exploratory Data Analysis",
        [
            "The EDA confirms that the dataset is large and balanced across major operational dimensions. Priority is distributed almost equally across Medium, Critical, Low, and High categories. Status is also distributed across Resolved, In Progress, Closed, New, and Open tickets.",
            "The topic distribution shows five major categories used by the project: General Inquiry, Network Issue, Hardware Failure, Access Request, and Software Bug. These categories later become KEDB categories when resolved and closed tickets are grouped by Topic.",
            "Source distribution is also balanced across Chat, Portal, Phone, and Email. Agent Group distribution covers Development, Network Ops, Security, Customer Service, and IT Support.",
        ],
        table={
            "headers": ["EDA Dimension", "Top Observations"],
            "rows": [
                ["Priority", "Medium 25,117; Critical 25,045; Low 25,014; High 24,824"],
                ["Status", "Resolved 20,134; In Progress 20,123; Closed 20,015; New 20,014; Open 19,714"],
                ["Topic", "General Inquiry 20,254; Network Issue 20,053; Hardware Failure 20,027; Access Request 19,923; Software Bug 19,743"],
                ["Source", "Chat 25,140; Portal 25,025; Phone 24,972; Email 24,863"],
                ["Agent Group", "Development 20,158; Network Ops 20,144; Security 19,985; Customer Service 19,884; IT Support 19,829"],
            ],
            "widths": [1.5, 4.7],
            "caption": "Table 3.2 Key EDA findings from the ticket dataset",
        },
    )
    section_page(
        doc,
        "3.3 KEDB Category Identification",
        [
            "The Known Error Database is identified and separated using the ticket Topic field. The training script filters tickets whose Status is Resolved or Closed, then groups them by Topic. Each unique topic becomes a known issue category.",
            "For each topic group, the project calculates occurrence count, typical agent group, average resolution time, solution suggestion, resolution steps, and auto-resolvable flag. A topic is marked auto-resolvable when it has sufficient historical occurrences, which is configured as 10 or more resolved/closed examples in the training script.",
            "This means the KEDB is not manually written from scratch. It is generated from historical incident outcomes, and its categories align with the major issue topics found during EDA.",
        ],
        table={
            "headers": ["KEDB Category", "Occurrences", "Typical Agent Group", "Solution Summary"],
            "rows": [
                ["General Inquiry", "8,130", "Security", "Provide documentation, FAQ resources, and escalate if needed."],
                ["Network Issue", "8,124", "Network Ops", "Check connectivity, DNS, adapter, and firewall rules."],
                ["Access Request", "8,005", "Network Ops", "Verify identity and grant permissions using least privilege."],
                ["Hardware Failure", "8,023", "IT Support", "Run diagnostics and verify physical connections."],
                ["Software Bug", "7,867", "Security", "Clear cache, reinstall software, and apply patches."],
            ],
            "widths": [1.45, 0.85, 1.35, 2.55],
            "caption": "Table 3.3 KEDB categories generated from resolved and closed tickets",
        },
    )
    section_page(
        doc,
        "3.4 Database Representation",
        [
            "ResolveNow AI stores operational tickets in SQLite using SQLAlchemy models. The Ticket table contains a database ID, ticket ID, topic, description, priority, status, creator, user email, assignment fields, created time, and AI result JSON.",
            "The use of SQLite makes the prototype lightweight and easy to run locally while still providing persistence. The dashboard and API can retrieve all tickets, filter by status or priority, update ticket status, close tickets, and compute analytics from the stored records.",
            "The AI result is stored as a JSON field so that the system can preserve the model output, action taken, KEDB match, solution suggestion, and automation decision along with the ticket record.",
        ],
        table={
            "headers": ["Field", "Purpose"],
            "rows": [
                ["ticket_id", "Human-readable incident number such as TCKT-100050"],
                ["topic / description", "Issue category and user-provided detail"],
                ["priority / status", "Routing and lifecycle control"],
                ["assigned_to / assigned_group", "Agent ownership information"],
                ["created_time", "Timestamp for dashboard and SLA calculations"],
                ["ai_result", "JSON record of prediction, action, KEDB match, and message"],
            ],
            "widths": [2.0, 4.2],
            "caption": "Table 3.4 SQLite ticket representation",
        },
    )

    section_page(
        doc,
        "CHAPTER 4 - METHODOLOGY",
        [
            "The methodology follows a staged architecture: input and validation, AI analysis, decision engine, database integration, communication workflow, dashboard reporting, and voice chatbot assistance. Each stage is implemented as a working part of the ResolveNow AI system.",
            "The dashboard captures user input. FastAPI validates and processes the request. The ML layer predicts the ticket attributes. The KEDB/RAG layer checks known issue knowledge. The decision engine determines whether to email the user, notify an agent, or assign the ticket. SQLite stores the result. The dashboard refreshes the view and analytics.",
        ],
        callout=("Methodology summary", "ResolveNow AI is a workflow automation system where AI output is immediately connected to operational action."),
    )
    section_page(
        doc,
        "4.1 System Architecture",
        [
            "The system is organized into a frontend layer, backend API layer, model and knowledge layer, persistence layer, and communication layer. This separation makes the project easier to understand and maintain.",
            "The frontend is implemented in a single dashboard HTML file with CSS and JavaScript. It communicates with FastAPI using fetch calls. The backend loads models during startup and exposes endpoints. The model layer includes pickled classifiers, encoders, TF-IDF vectorizer, FAISS index, KEDB JSON, and GraphSAGE prototype files.",
            "The communication layer includes SMTP email functions for solution emails, high priority alerts, escalation emails, and SLA warnings. The voice layer accepts recorded audio from the browser and forwards it to Sarvam speech-to-text through a backend proxy so the API key remains server-side.",
        ],
        table={
            "headers": ["Layer", "Main Components"],
            "rows": [
                ["Frontend", "Dashboard, ticket form, tables, analytics, KEDB view, voice bot"],
                ["Backend", "FastAPI routes, validation models, decision engine, SLA APIs"],
                ["AI/ML", "Priority model, topic model, TF-IDF, encoders, RAG/FAISS, GNN prototype"],
                ["Persistence", "SQLite database with SQLAlchemy Ticket model"],
                ["Communication", "Solution email, agent alert, escalation, SLA warning, confirmation links"],
            ],
            "widths": [1.5, 4.7],
            "caption": "Table 4.1 ResolveNow AI architecture layers",
        },
    )
    section_page(
        doc,
        "4.2 FastAPI Workflow",
        [
            "FastAPI acts as the central execution layer. It serves the dashboard, loads trained artifacts, validates new ticket payloads, creates ticket IDs, performs predictions, triggers the decision engine, sends emails, writes to the database, and returns JSON responses for the dashboard.",
            "The ticket creation endpoint is the main workflow entry point. It receives a NewTicket request, predicts topic and priority, applies user priority override if provided, calculates SLA deadlines, checks the KEDB, executes decision logic, sends appropriate emails, stores the ticket, and returns both ticket data and AI analysis.",
            "The same backend also supports health checks, ticket retrieval, manual resolve/close actions, KEDB lookup, analytics summary, prediction-only queries, SLA breached/warning list, confirmation links, agent resolution links, and voice ticket status lookup.",
        ],
        bullets=[
            "GET /dashboard serves the dashboard UI.",
            "POST /tickets/create creates and analyzes tickets.",
            "GET /tickets and GET /tickets/{ticket_id} support dashboard retrieval.",
            "PUT /tickets/{ticket_id}/resolve and DELETE /tickets/{ticket_id}/close support lifecycle actions.",
            "POST /voice/ticket-status supports Sarvam STT based ticket tracking.",
        ],
    )
    section_page(
        doc,
        "4.3 Model Pipeline",
        [
            "The working model pipeline uses TF-IDF text features and encoded categorical features to train two classifiers. The priority classifier uses RandomForestClassifier, and the topic classifier uses GradientBoostingClassifier. Both models are trained from historical ticket data and saved as pickle artifacts.",
            "The input text combines Topic, Source, Product group, and Support Level. Numeric features include encoded Source, Product group, and Agent Group. These are horizontally stacked into a single feature matrix for model training and prediction.",
            "During runtime, FastAPI loads the pickled models and encoders. For each ticket, it transforms the text into TF-IDF, encodes categorical values safely, predicts priority and topic, and returns confidence values.",
        ],
        table={
            "headers": ["Model Artifact", "Role"],
            "rows": [
                ["priority_model.pkl", "Predicts ticket priority class"],
                ["topic_model.pkl", "Predicts issue topic class"],
                ["tfidf.pkl", "Transforms combined ticket text into vector features"],
                ["priority_enc.pkl / topic_enc.pkl", "Converts numeric labels back to readable classes"],
                ["source_enc.pkl / product_enc.pkl / agent_enc.pkl", "Encodes categorical operational context"],
            ],
            "widths": [2.3, 3.9],
            "caption": "Table 4.2 Runtime model artifacts",
        },
    )
    section_page(
        doc,
        "4.4 BERT and Semantic Understanding",
        [
            "The final presentation describes BERT for topic classification using semantic meaning. In the project codebase, semantic representation appears in the RAG pipeline through SentenceTransformer using all-MiniLM-L6-v2, a compact transformer model commonly used for sentence embeddings.",
            "Instead of matching only exact words, embeddings allow the system to represent the meaning of ticket text and KEDB entries numerically. Similar issues can be close in vector space even when the wording is different.",
            "For the deployed classifier, the working code uses TF-IDF and encoded features. For presentation purposes, it is accurate to explain that BERT-style semantic embeddings support the knowledge retrieval side of the project, while the classification artifact uses a traditional ML implementation trained on ticket features.",
        ],
    )
    section_page(
        doc,
        "4.5 RAG and FAISS Retrieval",
        [
            "RAG in ResolveNow AI means retrieval-augmented solution assistance. KEDB entries are converted into embeddings and stored in a FAISS index. When a query is created from ticket fields, the same embedding model converts the query into a vector, and FAISS searches for the nearest KEDB entries.",
            "The RAG module uses SentenceTransformer to encode the query and FAISS IndexFlatL2 to retrieve similar known issues. The retrieved result can then be used as a solution suggestion or context for a generated resolution plan.",
            "In the current working ticket flow, direct KEDB matching by topic is used inside the FastAPI decision engine. The RAG module is available as a knowledge retrieval component and is used by supporting agent/ticket handler scripts.",
        ],
        bullets=[
            "Embedding model: SentenceTransformer all-MiniLM-L6-v2.",
            "Vector store: FAISS index stored under models/faiss_index/kedb.index.",
            "Knowledge source: KEDB JSON entries built from resolved and closed tickets.",
            "Output: ranked known issue candidates and solution suggestions.",
        ],
    )
    section_page(
        doc,
        "4.6 GNN and GraphSAGE",
        [
            "The GNN prototype uses Neo4j and PyTorch Geometric. Tickets and resolutions are represented as graph nodes, and relationships such as Ticket resolved by Resolution are modeled as graph edges. This graph is exported into a PyTorch Geometric Data object.",
            "GraphSAGE is used because it can learn from a node's neighborhood. For an ITSM system, this means a ticket can be understood not only from its own priority but also from connected resolution patterns and similar historical cases.",
            "The prototype defines a two-layer GraphSAGE model and trains it to predict SLA labels as either SLA Met or SLA Breach Risk. This demonstrates how graph reasoning can extend the system beyond keyword retrieval.",
        ],
        table={
            "headers": ["GNN Element", "Project Meaning"],
            "rows": [
                ["Ticket node", "Incident record with priority, topic, and SLA label"],
                ["Resolution node", "Historical resolution text or known fix"],
                ["RESOLVED_BY edge", "Relationship between a ticket and its resolution"],
                ["GraphSAGE", "Learns from ticket neighborhoods to predict SLA outcome"],
                ["Prediction output", "SLA Met or SLA Breach Risk"],
            ],
            "widths": [1.8, 4.4],
            "caption": "Table 4.3 GNN representation in ResolveNow AI",
        },
    )
    section_page(
        doc,
        "4.7 SHAP Explainability",
        [
            "SHAP is used to explain which features influence model decisions. In the training script, SHAP TreeExplainer is applied to the topic classifier to generate a feature importance plot. This helps identify which words and encoded fields drive topic classification.",
            "In the RAG module, SHAP is also included conceptually for explaining query embeddings. It highlights which parts of a query contribute to retrieval behavior. This makes the retrieval process less of a black box.",
            "For the report, SHAP is described as an explainability layer. It supports trust by showing why a classifier or retrieval model is influenced by certain ticket words or fields.",
        ],
        image={"path": "models/shap_topic_importance.png", "caption": "Figure 4.1 SHAP feature importance generated by the training workflow", "width": 5.7},
    )

    section_page(
        doc,
        "CHAPTER 5 - CORE STRATEGY AND IMPLEMENTATION",
        [
            "The implementation strategy is to connect AI analysis with immediate operational action. A prediction alone does not solve ITSM delay. ResolveNow AI therefore uses prediction, retrieval, decision rules, email, database persistence, and dashboard refresh as one continuous process.",
            "The core strategy separates tickets into high/critical and low/medium paths. Known high/critical issues are escalated to a human agent with a suggested solution. Unknown high/critical issues are escalated directly. Known low/medium issues trigger a solution email to the user and wait for confirmation. Unknown low/medium issues are assigned to an agent.",
        ],
    )
    section_page(
        doc,
        "5.1 Decision Engine",
        [
            "The decision engine is implemented in the process_ticket function. It checks the predicted priority and the KEDB match. Based on those two dimensions, it returns an action such as SUGGESTION_ONLY, ESCALATED_TO_HUMAN, EMAIL_SENT_PENDING_CUSTOMER, or PREDICTED_AND_ASSIGNED.",
            "The action result includes the ticket ID, priority, known issue flag, solution suggestion, resolution steps, auto resolution flags, required human flag, KEDB match, email requirement, and a message. This structured result is stored with the ticket and displayed on the dashboard.",
        ],
        table={
            "headers": ["Condition", "System Action"],
            "rows": [
                ["High/Critical + Known Issue", "Assign agent and send solution suggestion to assignee"],
                ["High/Critical + Unknown Issue", "Escalate to human agent without auto-resolution"],
                ["Low/Medium + Known Issue", "Email solution to user and mark Pending Customer"],
                ["Low/Medium + Unknown Issue", "Assign to support group and mark In Progress"],
            ],
            "widths": [2.4, 3.8],
            "caption": "Table 5.1 Priority-aware decision logic",
        },
    )
    section_page(
        doc,
        "5.2 Dashboard Implementation",
        [
            "The dashboard is implemented as a professional single-page frontend with multiple views. It includes Dashboard, All Tickets, New Ticket, KEDB, and Analytics pages. JavaScript fetch calls connect the dashboard to FastAPI endpoints.",
            "The Dashboard page shows total tickets, auto-resolved count, human-needed count, automation rate, ticket pipeline status, recent tickets, and SLA status. The All Tickets page supports search, filtering, and detail viewing. The New Ticket page allows inline ticket creation and shows the AI result card after submission.",
            "The KEDB page displays known issues and solution suggestions. The Analytics page summarizes ticket performance. A report download button exports ticket data as CSV, and the floating voice bot allows hands-free ticket status lookup.",
        ],
        bullets=[
            "Live metric cards and ticket pipeline.",
            "Searchable and filterable ticket table.",
            "Ticket detail side panel with AI analysis.",
            "KEDB display for known issue solutions.",
            "Analytics summary and CSV report download.",
            "Bottom-right voice chatbot for ticket status.",
        ],
    )
    section_page(
        doc,
        "5.3 Email Automation",
        [
            "Email automation is one of the most important working outputs. The backend email service contains functions for solution emails, high priority emails, escalation emails, and SLA warning emails. These emails transform model decisions into user or agent communication.",
            "For low and medium known issues, the user receives a solution email with steps and confirmation links. If the user confirms that the issue is solved, the ticket is closed. If the user indicates that the issue is still present, the ticket is escalated to an agent.",
            "For high and critical tickets, the agent receives an urgent alert and, when available, a solution suggestion. The user can also receive acknowledgement that the issue has been assigned to a human agent.",
        ],
        table={
            "headers": ["Email Type", "Trigger", "Recipient"],
            "rows": [
                ["Solution email", "Low/Medium known issue", "User"],
                ["High priority alert", "High/Critical ticket", "Agent"],
                ["Escalation email", "User says issue not resolved", "Agent"],
                ["SLA warning email", "75 percent SLA time consumed", "Agent or support team"],
            ],
            "widths": [1.7, 2.7, 1.8],
            "caption": "Table 5.2 Email automation workflow",
        },
    )
    section_page(
        doc,
        "5.4 SLA Monitoring",
        [
            "SLA monitoring is implemented through backend/sla_service.py. The service defines SLA response and restoration windows by priority. Critical and High tickets have two-hour response and twenty-four-hour restoration windows, while Medium and Low tickets have twenty-four-hour response and ninety-six-hour restoration windows.",
            "The backend calculates response and restoration deadlines at ticket creation time. It can recalculate live SLA status, determine remaining time, identify breaches, and identify tickets that have consumed 75 percent of their SLA time.",
            "The dashboard calls the SLA breached endpoint and shows breached and warning tickets. This makes SLA monitoring part of the live operating view rather than a separate offline report.",
        ],
        table={
            "headers": ["Priority", "Response SLA", "Restoration SLA"],
            "rows": [
                ["Critical", "2 hours", "24 hours"],
                ["High", "2 hours", "24 hours"],
                ["Medium", "24 hours", "96 hours"],
                ["Low", "24 hours", "96 hours"],
            ],
            "widths": [1.7, 2.2, 2.3],
            "caption": "Table 5.3 SLA configuration used by ResolveNow AI",
        },
    )
    section_page(
        doc,
        "5.5 Voice Chatbot with Sarvam API",
        [
            "The voice chatbot is implemented as a small bottom-right robot widget in the dashboard. The user opens the widget, clicks Start Voice, and speaks the incident number. The browser records the audio using MediaRecorder and sends it to the backend.",
            "The backend receives the audio at POST /voice/ticket-status, forwards it to Sarvam speech-to-text using the configured API key, extracts the ticket number from the transcript, retrieves the matching ticket, and returns the status, priority, topic, assignment, and a spoken response.",
            "The frontend displays the transcript and status response and uses browser speech synthesis to read the answer aloud. This feature gives users a faster and more accessible way to track ticket status.",
        ],
        callout=("Chatbot statement", "The ResolveNow AI chatbot captures a spoken incident number, converts speech to text using Sarvam API, finds the matching ticket, and reads the current status back to the user."),
    )
    section_page(
        doc,
        "5.6 API Endpoint Summary",
        [
            "The backend provides a set of REST endpoints that support the dashboard and automated workflows. The final presentation mentions 17 working REST API endpoints. The current FastAPI implementation includes endpoints for dashboard serving, health, ticket creation, voice status, ticket confirmation, agent resolution, ticket listing, ticket detail, resolve, close, KEDB retrieval, analytics, prediction, ticket SLA, and breached SLA monitoring.",
            "These APIs are important because they make the system modular. The dashboard can be replaced by another frontend in the future, and the same backend intelligence can be reused.",
        ],
        table={
            "headers": ["Endpoint Group", "Purpose"],
            "rows": [
                ["Dashboard and health", "Serve UI and verify model readiness"],
                ["Ticket creation", "Create, predict, route, email, and store ticket"],
                ["Ticket management", "List, view, resolve, and close tickets"],
                ["KEDB and prediction", "Retrieve known errors and run prediction-only checks"],
                ["SLA", "Calculate breached and warning tickets"],
                ["Voice chatbot", "Transcribe incident number and return ticket status"],
            ],
            "widths": [2.0, 4.2],
            "caption": "Table 5.4 Backend API capability summary",
        },
    )

    section_page(
        doc,
        "CHAPTER 6 - ACTIONABLE OUTCOMES",
        [
            "The project outcome is a functional prototype that demonstrates end-to-end ITSM automation. The application is accessible through the local dashboard and supports the complete ticket lifecycle from creation to analysis, routing, communication, storage, and status monitoring.",
            "The working portions include ticket creation, model loading, prediction, KEDB matching, priority-based decision logic, email workflow, SQLite storage, dashboard display, analytics summary, SLA warning view, CSV report download, and voice chatbot status query.",
        ],
    )
    section_page(
        doc,
        "6.1 Functional Outputs",
        [
            "The final output described in the presentation is a deployed local autonomous ITSM platform available at localhost:8000/dashboard. The report describes this as the working prototype interface because it depends on running the FastAPI server locally.",
            "A typical flow is: ticket created, model predicts topic and priority, KEDB identifies known issue, decision engine routes by priority, emails are sent, ticket is stored in SQLite, dashboard refreshes metrics, and the user can confirm resolution or check status through the chatbot.",
        ],
        bullets=[
            "Live dashboard with five functional pages.",
            "Real-time AI analysis result after ticket submission.",
            "Known issue matching through generated KEDB.",
            "Automated email communication with confirmation buttons.",
            "SQLite database audit trail.",
            "Voice chatbot for incident status tracking.",
        ],
    )
    section_page(
        doc,
        "6.2 Model Evaluation Outputs",
        [
            "The training script prints priority model accuracy and topic model accuracy and generates several visual evaluation artifacts. The presentation reports BERT topic classification accuracy of 94.2 percent with F1 of 0.93, RAG retrieval accuracy of 89 percent on 20 test queries, and SHAP explaining top features.",
            "In the working code, the topic classifier is evaluated using accuracy, precision, recall, F1, classification report, confusion matrix, overall metrics, and SHAP feature importance. These outputs are saved under the models directory.",
            "The report separates conceptual presentation metrics from working artifact evidence. This is useful for academic clarity because the prototype combines both deployed ML artifacts and experimental AI modules.",
        ],
        image={"path": "models/topic_overall_metrics.png", "caption": "Figure 6.1 Overall topic classification performance chart", "width": 5.8},
    )
    section_page(
        doc,
        "6.3 Confusion Matrix and Class-Level Performance",
        [
            "The confusion matrix helps understand where the topic classifier performs well and where classes may be confused. This is important because topic prediction influences KEDB lookup, routing, and solution suggestion.",
            "The precision, recall, and F1 chart provides class-level evaluation. Instead of relying only on a single accuracy score, the project records how each issue category behaves. This is especially useful in ITSM because an error in a high-risk issue category can affect SLA outcomes.",
            "The generated heatmap and bar chart provide visual evidence for model quality and make the report easier to evaluate.",
        ],
        image={"path": "models/topic_confusion_matrix.png", "caption": "Figure 6.2 Topic classification confusion matrix", "width": 5.8},
    )
    section_page(
        doc,
        "6.4 Business Impact",
        [
            "The final presentation estimates business impact by comparing manual ITSM operations with ResolveNow AI. The key improvement areas are MTTR reduction, support team resource reduction, SLA breach reduction, automation rate, and misrouting reduction.",
            "These numbers are presented as approximate business impact estimates for the prototype. They help demonstrate how the system can create value if deployed into an enterprise ITSM environment.",
        ],
        table={
            "headers": ["Metric", "Before ResolveNow AI", "After ResolveNow AI", "Improvement"],
            "rows": [
                ["MTTR", "48-72 hours", "2-4 hours", "85-90 percent reduction"],
                ["L1 support team", "6 members", "3 members", "50 percent resource cost reduction"],
                ["SLA breach rate", "About 28 percent", "About 6 percent", "78 percent penalty reduction"],
                ["Auto-resolved tickets", "0 percent", "About 65 percent of L/M known volume", "65 percent automation rate"],
                ["Misrouting rate", "About 22 percent", "Less than 2 percent", "91 percent improvement"],
            ],
            "widths": [1.35, 1.55, 1.65, 1.65],
            "caption": "Table 6.1 Estimated business impact from the final presentation",
        },
    )
    section_page(
        doc,
        "6.5 User and Agent Outcomes",
        [
            "For users, the main benefit is faster communication. Instead of waiting for manual triage, users can receive a solution email for known low/medium issues and confirm whether it solved the problem. They can also check ticket status through the dashboard or voice bot.",
            "For agents, the main benefit is better focus. High and critical tickets are prioritized, and known issue context is included in the alert. Agents spend less time searching for the first solution and more time resolving genuinely complex issues.",
            "For managers, the dashboard provides visibility. Ticket counts, automation rate, human-required load, KEDB size, SLA warnings, and analytics summary help evaluate operational performance.",
        ],
    )

    section_page(
        doc,
        "CHAPTER 7 - CHALLENGES",
        [
            "The project faced technical and design challenges typical of applied AI systems. The first challenge was integrating multiple components - data preparation, model training, vector retrieval, graph reasoning, backend APIs, frontend dashboard, email, database, and voice chatbot - into a single coherent workflow.",
            "The second challenge was ensuring that model outputs led to meaningful actions. A classifier score alone is not useful unless the system can decide what to do next. The decision engine therefore became a critical bridge between AI and operations.",
        ],
    )
    section_page(
        doc,
        "7.1 Data and Modeling Challenges",
        [
            "The dataset is large and balanced, but real ITSM data can contain noisy descriptions, missing fields, inconsistent categories, duplicated tickets, and changing issue patterns. The project handles missing values and constructs combined text features, but production deployment would require stronger data quality controls.",
            "Another modeling challenge is explaining the difference between conceptual BERT/RAG/GNN architecture and the exact working prototype. The report addresses this by clearly describing which components are deployed directly in FastAPI and which components are prototype or supporting modules.",
            "Model evaluation must also be interpreted carefully. Classification accuracy applies to supervised labels, while RAG requires retrieval accuracy or relevance evaluation, and GNN requires comparison between predicted and actual SLA labels.",
        ],
    )
    section_page(
        doc,
        "7.2 Integration Challenges",
        [
            "Integration required careful handling of asynchronous workflows. A ticket may be created, emailed, stored, updated by user confirmation, and later shown on the dashboard. Keeping these states consistent requires the database to act as the source of truth.",
            "The email workflow also required confirmation URLs. When the user clicks Solved or Still Having Issue, the backend changes ticket status accordingly. This makes the email not just a notification but part of the workflow.",
            "The voice chatbot introduced browser permission, audio encoding, API key security, Sarvam model versioning, and transcript parsing challenges. The solution keeps the API key on the backend and extracts ticket numbers from the transcribed text.",
        ],
    )
    section_page(
        doc,
        "7.3 Deployment and Operational Challenges",
        [
            "The current deployment is local through Uvicorn and FastAPI. This is appropriate for a capstone prototype, but production deployment would need secure hosting, environment variable management, HTTPS, authentication, logging, monitoring, and database migration support.",
            "The system also depends on model and KEDB artifacts being present at startup. A production-grade implementation would include artifact versioning, health checks for model readiness, backup strategies, and scheduled retraining.",
            "Another operational challenge is that automation must be controlled carefully. The project routes high and critical tickets to humans instead of auto-closing them, which is an important design choice for safety and accountability.",
        ],
    )

    section_page(
        doc,
        "CHAPTER 8 - CONCLUSIONS AND RECOMMENDATIONS",
        [
            "ResolveNow AI successfully demonstrates that autonomous ITSM is achievable using a combination of machine learning, KEDB construction, retrieval, graph reasoning, rule-based decision logic, backend automation, dashboard visibility, email workflow, and voice assistance.",
            "The project converts historical ticket data into working intelligence. It predicts ticket attributes, identifies known issue categories, retrieves or attaches solution knowledge, routes tickets by priority, communicates with users and agents, monitors SLA status, and stores the lifecycle in a database.",
            "The most important conclusion is that AI value appears when prediction is connected to action. ResolveNow AI does this by turning model output into email, assignment, dashboard state, and ticket status changes.",
        ],
        callout=("Conclusion", "The prototype proves the feasibility of an AI-assisted ITSM workflow that reduces manual effort while preserving human control for high-priority incidents."),
    )
    section_page(
        doc,
        "8.1 Recommendations",
        [
            "The project can be extended into a stronger enterprise system through integration with real ticketing tools, continuous learning, stronger authentication, and production deployment practices. The most valuable next step is to connect ResolveNow AI with an ITSM platform such as ServiceNow or Jira through APIs or MCP-based agents.",
            "A self-learning KEDB should be implemented so every resolved ticket can enrich the knowledge base after validation. When new categories or resolution patterns reach a threshold, the system can trigger retraining for the classifier and update FAISS/GNN artifacts.",
            "Future work should also improve multilingual support, chatbot conversation depth, explainability dashboards, and governance rules for when automation is allowed.",
        ],
        bullets=[
            "Integrate with enterprise ticketing tools using APIs or MCP agents.",
            "Build a continuous retraining pipeline for classifier and retrieval artifacts.",
            "Add authentication, authorization, and role-based dashboard views.",
            "Add production logging, monitoring, and audit reports.",
            "Expand Sarvam voice support into a multilingual conversational assistant.",
        ],
    )
    section_page(
        doc,
        "8.2 Final Summary",
        [
            "ResolveNow AI is a complete capstone prototype that combines analytics, AI, backend engineering, frontend design, and workflow automation. It moves beyond static prediction by implementing a working ticket lifecycle.",
            "The final system shows how data science can be applied to a business operations problem. It reduces repetitive search, supports faster resolution, improves routing consistency, and gives users and managers better visibility.",
            "The project is therefore positioned as an autonomous ITSM incident intelligence system with practical value and clear future extension paths.",
        ],
    )

    section_page(
        doc,
        "APPENDIX A - SAMPLE API AND DATA STRUCTURES",
        [
            "This appendix summarizes important implementation structures used by the project. The NewTicket request model includes ticket topic, source, product group, support level, agent group, description, country, creator, user priority, and user email.",
            "The Ticket database model stores ticket ID, topic, description, priority, status, creator, user email, assignment, created time, and AI result JSON. This structure is sufficient for dashboard display and audit.",
        ],
        table={
            "headers": ["Structure", "Fields"],
            "rows": [
                ["NewTicket", "topic, source, product_group, support_level, agent_group, description, country, created_by, user_priority, user_email"],
                ["Ticket", "ticket_id, topic, description, priority, status, created_by, user_email, assigned_to, assigned_group, created_time, ai_result"],
                ["AI Result", "priority, known issue, action taken, KEDB match, solution suggestion, human required flag, email required flag"],
                ["Voice Status Response", "transcript, ticket ID, status, priority, topic, assigned_to, assigned_group, spoken_response"],
            ],
            "widths": [1.6, 4.6],
            "caption": "Table A.1 Key API and database structures",
        },
    )
    section_page(
        doc,
        "APPENDIX B - DEMONSTRATION SCRIPT",
        [
            "The demonstration can begin by running the FastAPI server using uvicorn backend.api:app --reload --port 8000. After the server loads models, the dashboard can be opened at http://127.0.0.1:8000/dashboard.",
            "The first demo step is to create a low or medium known ticket such as Network Issue. The AI result card should show known issue detection and email pending customer action. The KEDB page can then be opened to show the corresponding known solution.",
            "The second demo step is to create a high or critical ticket. The result should show human escalation or suggestion to the agent. The All Tickets page can be used to open the detail panel, and the Analytics page can be used to show summary metrics.",
            "The third demo step is to open the voice bot and speak an incident number such as TCKT-100050. The chatbot should send audio to Sarvam speech-to-text, capture the incident number, retrieve ticket status, and speak the result.",
        ],
    )
    section_page(
        doc,
        "APPENDIX C - PRESENTATION ALIGNMENT",
        [
            "This report is aligned with the final presentation titled ResolveNow AI - Autonomous ITSM Incident Intelligence System. The presentation introduced the problem, data source, EDA, methodology, model evaluation, results, business impact, final output, implementation, deployment, conclusion, future work, and literature survey.",
            "The report expands those slide points into a formal academic structure. It also includes additional explanation of how KEDB categories are identified, how SHAP supports explainability, how BERT/SentenceTransformer embeddings support semantic retrieval, how the GNN prototype works, and how the FastAPI/dashboard implementation supports the workflow.",
            "The report intentionally focuses on the portions that are working in the project repository: dashboard, FastAPI endpoints, ticket creation, model loading, KEDB matching, email automation, database storage, SLA monitoring, report download, and Sarvam voice chatbot.",
        ],
    )

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_report()
    print(path.resolve())
